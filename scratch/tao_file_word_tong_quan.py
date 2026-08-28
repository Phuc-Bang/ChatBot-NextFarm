#!/usr/bin/env python3
"""
tao_file_word_tong_quan.py — Sinh file Word (.docx) tong quan cau truc du an NextFarm AI.
"""

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from pathlib import Path

OUT_DOCX = Path("d:/NextFarm/docs/TONG_QUAN_CAU_TRUC_DU_AN_NEXTFARM.docx")
ROOT_DOCX = Path("d:/NextFarm/TONG_QUAN_CAU_TRUC_DU_AN_NEXTFARM.docx")

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def format_row(row, texts, is_header=False, fill_hex="16A34A", is_even=False):
    for i, text in enumerate(texts):
        cell = row.cells[i]
        cell.text = text
        fill = fill_hex if is_header else ("F8FAFC" if is_even else "FFFFFF")
        set_cell_background(cell, fill)
        set_cell_margins(cell, top=120 if is_header else 90, bottom=120 if is_header else 90, left=140, right=140)
        p = cell.paragraphs[0]
        for r in p.runs:
            r.font.name = "Arial"
            if is_header:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(10)
            else:
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(30, 41, 59)

def main():
    doc = docx.Document()

    # Page Margins
    for s in doc.sections:
        s.top_margin = Inches(0.75)
        s.bottom_margin = Inches(0.75)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)

    # Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(30, 41, 59)
    normal_style.paragraph_format.line_spacing = 1.25
    normal_style.paragraph_format.space_after = Pt(5)

    # ==================== COVER / HEADER ====================
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(3)
    run_title = title_p.add_run("TỔNG QUAN CẤU TRÚC DỰ ÁN NEXTFARM AI")
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(22, 163, 74) # Primary Green

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(12)
    run_sub = sub_p.add_run("Bài Toán A: Chống Bịa Đặt & AI Safety trong Trợ Lý Nông Nghiệp Thông Minh")
    run_sub.font.size = Pt(11.5)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGBColor(71, 85, 105)

    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False

    format_row(meta_table.rows[0], ["• Dự án: NextFarm AI (PoC Giai đoạn 1 & 2)", "• Quy chuẩn kỹ thuật: v2.0 Standard"], fill_hex="F1F5F9")
    format_row(meta_table.rows[1], ["• Thời gian: Tháng 08/2026", "• Trạng thái: Hoàn tất 100% · 367 Tests Xanh"], fill_hex="F1F5F9")

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ==================== PHẦN 1: MỤC TIÊU CỐT LÕI ====================
    h1 = doc.add_heading("1. Mục Tiêu Cốt Lõi & Giải Pháp 4 Hiện Tượng Rủi Ro", level=1)
    h1.paragraph_format.space_before = Pt(10)
    h1.paragraph_format.space_after = Pt(4)
    for r in h1.runs:
        r.font.color.rgb = RGBColor(15, 23, 42)
        r.font.size = Pt(13)
        r.font.bold = True

    p = doc.add_paragraph()
    p.add_run("Hệ thống giải quyết triệt để ").font.color.rgb = RGBColor(30, 41, 59)
    p.add_run("4 hiện tượng rủi ro nông học ").bold = True
    p.add_run("của LLM trần khi tư vấn cho 3 cây trồng: ")
    p.add_run("Lúa, Cà chua và Dưa chuột:").bold = True

    t_risk = doc.add_table(rows=5, cols=3)
    t_risk.alignment = WD_TABLE_ALIGNMENT.CENTER
    format_row(t_risk.rows[0], ["Hiện Tượng", "Hành Vi LLM Trần (Nguy Hiểm)", "Giải Pháp NextFarm PoC"], is_header=True, fill_hex="0F172A")

    risk_data = [
        ("A1: Bịa số liệu vườn", "Nông dân hỏi độ ẩm đất hiện tại; LLM tự bịa 75%, 30°C dù không có cảm biến.", "Intent Router phát hiện 'garden_data', từ chối ngay trong 1ms (0 token)."),
        ("A2: Bịa tính năng app", "Nông dân hỏi tính năng dự báo giá nông sản; LLM tự nhận có.", "Intent Router nhận diện 'product_feature', từ chối lịch sự minh bạch."),
        ("A3: Sai cây / Vùng miền", "Tư vấn cây ngoài danh mục (sầu riêng, cam) hoặc áp sai thời vụ miền Bắc cho miền Tây.", "Scope Check giới hạn cứng 3 cây; RRF cộng điểm ưu tiên vùng miền."),
        ("A4: Tiếng Việt không dấu", "Nông dân gõ 'ca chua bi sau', LLM đoán sai dấu hoặc hiểu nhầm từ viết tắt.", "Chuẩn hoá 4 lớp + Trigram matching ở CSDL, không ép LLM đoán dấu.")
    ]

    for idx, r in enumerate(risk_data, start=1):
        format_row(t_risk.rows[idx], [r[0], r[1], r[2]], is_even=(idx % 2 == 0))
        t_risk.rows[idx].cells[0].paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ==================== PHẦN 2: CẤU TRÚC CÂY THƯ MỤC ====================
    h2 = doc.add_heading("2. Cấu Trúc Cây Thư Mục & Phân Bổ Chức Năng", level=1)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(4)
    for r in h2.runs:
        r.font.color.rgb = RGBColor(15, 23, 42)
        r.font.size = Pt(13)
        r.font.bold = True

    t_tree = doc.add_table(rows=8, cols=3)
    t_tree.alignment = WD_TABLE_ALIGNMENT.CENTER
    format_row(t_tree.rows[0], ["Thư Mục / Tệp", "Vai Trò Chức Năng", "Thành Phần Trọng Tâm"], is_header=True, fill_hex="16A34A")

    tree_data = [
        ("app/", "Mã nguồn lõi ứng dụng FastAPI", "main.py, services/ (pipeline, intent, normalization, retrieval, llm, grounding)"),
        ("crawler/", "Thu thập tài liệu khuyến nông .gov.vn", "main.py, spiders/ (HTML + PDF), data/ (manifest.json, text/, raw/)"),
        ("knowledge/", "Kho tri thức & Kiểm duyệt 2 luồng", "review/ (documents.yaml, chunks.yaml, facts.yaml), lexicon/ (stopwords, từ điển), chunking/"),
        ("evaluation/", "Bộ đo lường & Tập kiểm thử đóng băng", "data/ (222 test cases v3, 30 regional cases), runners/ (C0, C1, C2), results/ (expert_scores.json), scripts/"),
        ("frontend/", "Giao diện Web đa người dùng", "chat.html (Nông dân), admin.html (Quản trị), report.html (Báo cáo), expert.html (Chấm điểm)"),
        ("docs/", "Tài liệu kỹ thuật & Bàn giao chính thức", "GIAO_HANG_NEXTFARM.md, BAO_CAO_TONG_KET_NEXTFARM.md, TRIEN_KHAI_DOCKER.md, reports/"),
        ("tests/", "Bộ kiểm thử tự động (367 tests)", "test_intent_router.py, test_normalization.py, test_keyword_retrieval.py, test_fewshot_router.py...")
    ]

    for idx, r in enumerate(tree_data, start=1):
        format_row(t_tree.rows[idx], [r[0], r[1], r[2]], is_even=(idx % 2 == 0))
        t_tree.rows[idx].cells[0].paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ==================== PHẦN 3: KIẾN TRÚC PIPELINE 7 CHẶNG ====================
    h3 = doc.add_heading("3. Kiến Trúc Pipeline 7 Chặng & 3 Tầng Guardrail", level=1)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(4)
    for r in h3.runs:
        r.font.color.rgb = RGBColor(15, 23, 42)
        r.font.size = Pt(13)
        r.font.bold = True

    steps = [
        ("Chặng 1: Chuẩn hoá tiếng Việt 4 lớp (0ms, 0 Token)", "Unicode NFC, hạ chữ thường, tra cứu từ điển từ đồng nghĩa nông học (sào Bắc Bộ vs sào Trung Bộ) và mở rộng viết tắt an toàn."),
        ("Chặng 2: Intent Router 2 tầng (Chặn sớm trong 5ms)", "Tầng 1: Rule-based siêu nhanh bắt chắc chắn các mẫu lệnh thiết bị, hỏi vườn -> Chặn ngay 141/222 case sớm. Tầng 2: LLMFewShotRouter (§11.3 & §40.2 Mục 9) xử lý câu hỏi biên."),
        ("Chặng 3: Scope Check & Bóc tách thực thể (4ms)", "Kiểm soát cứng 3 cây trồng: Lúa, Cà chua, Dưa chuột. Ngoài danh mục thì từ chối; thiếu chủ ngữ thì dùng template hỏi lại rõ ràng."),
        ("Chặng 4: Hybrid Retrieval & Cổng DEC-005", "Truy xuất 3 kênh: Vector (Local 768-dim) + Full-Text Search (Simple OR) + Trigram hợp nhất bằng RRF. Chỉ đọc từ view indexable_chunk đã được phê duyệt."),
        ("Chặng 5: Evidence Pack Generator", "Đóng gói ngữ cảnh JSON nguyên văn kèm chunk_id, tên cơ quan phát hành (Sở NN&PTNT Ninh Bình, Hà Tĩnh...) và Source Tier."),
        ("Chặng 6: LLM Generator (Gemini 3.1 Flash Lite)", "Đọc Evidence Pack, cấu hình thinking_budget=0, trả về câu trả lời có gắn mã trích dẫn nguồn [cid] trực tiếp cho từng khẳng định."),
        ("Chặng 7: Grounding Validator Đa Tầng", "Kiểm tra cấu trúc trích dẫn; bóc tách toàn bộ con số (pH, kg, %) đối soát với tài liệu gốc. Lệch số hoặc thiếu chứng cứ -> lập tức chuyển sang từ chối an toàn.")
    ]

    for st_title, st_desc in steps:
        p_st = doc.add_paragraph()
        p_st.paragraph_format.left_indent = Inches(0.2)
        p_st.paragraph_format.space_after = Pt(2)
        r_num = p_st.add_run(f"• {st_title}: ")
        r_num.bold = True
        r_num.font.color.rgb = RGBColor(22, 163, 74)
        p_st.add_run(st_desc).font.color.rgb = RGBColor(51, 65, 85)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ==================== PHẦN 4: KẾT QUẢ ĐỐI CHỨNG C0 - C1 - C2 ====================
    h4 = doc.add_heading("4. Bảng Kết Quả Thực Nghiệm Đối Chứng (C0 · C1 · C2)", level=1)
    h4.paragraph_format.space_before = Pt(12)
    h4.paragraph_format.space_after = Pt(4)
    for r in h4.runs:
        r.font.color.rgb = RGBColor(15, 23, 42)
        r.font.size = Pt(13)
        r.font.bold = True

    t_comp = doc.add_table(rows=8, cols=4)
    t_comp.alignment = WD_TABLE_ALIGNMENT.CENTER
    format_row(t_comp.rows[0], ["Chỉ Số Kiểm Định (222 Test Cases)", "C0: LLM Trần", "C1: RAG Thuần", "C2: NextFarm PoC"], is_header=True, fill_hex="2563EB")

    comp_rows = [
        ("Tổng số ca bịa số liệu nông học", "61 ca (27.5%)", "23 ca (10.4%)", "0 ca (0.0%)"),
        ("Tỷ lệ trả lời sai (False Answer Rate)", "77.0%", "23.0%", "0.9%"),
        ("Độ chính xác khi trả lời", "1.2%", "23.9%", "90.9%"),
        ("Tỷ lệ bắt ca cần từ chối (Recall)", "3.4%", "69.6%", "100.0% (148/148)"),
        ("Độ trễ trung vị p50", "2.621 ms", "2.555 ms", "15 ms (Giảm 99.4%)"),
        ("Chi phí API / 222 câu hỏi", "$0.0369", "$0.1231", "$0.0527 (~$0.0002/câu)"),
        ("Bộ kiểm thử tự động", "—", "—", "367 / 367 Tests Xanh")
    ]

    for idx, r in enumerate(comp_rows, start=1):
        format_row(t_comp.rows[idx], [r[0], r[1], r[2], r[3]], is_even=(idx % 2 == 0))
        t_comp.rows[idx].cells[0].paragraphs[0].runs[0].font.bold = True
        t_comp.rows[idx].cells[3].paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ==================== PHẦN 5: NHỮNG GÌ CHƯA LÀM & LỘ TRÌNH ====================
    h5 = doc.add_heading("5. Giới Hạn Đã Biết & Lộ Trình Mở Rộng", level=1)
    h5.paragraph_format.space_before = Pt(12)
    h5.paragraph_format.space_after = Pt(4)
    for r in h5.runs:
        r.font.color.rgb = RGBColor(15, 23, 42)
        r.font.size = Pt(13)
        r.font.bold = True

    t_road = doc.add_table(rows=4, cols=3)
    t_road.alignment = WD_TABLE_ALIGNMENT.CENTER
    format_row(t_road.rows[0], ["Hạng Mục", "Trạng Thái Hiện Tại (PoC)", "Kế Hoạch Giai Đoạn Tiếp Theo"], is_header=True, fill_hex="D97706")

    road_data = [
        ("Dữ liệu cảm biến IoT", "Chưa kết nối API IoT NextFarm; từ chối an toàn khi hỏi số liệu vườn.", "Giai đoạn 3: Kết nối API trạm đo thực tế theo từng vườn của người dùng."),
        ("Tài liệu sản phẩm NextFarm", "Chưa có tài liệu chính thức; từ chối câu hỏi giá/tính năng app.", "Giai đoạn 2: Bổ sung cẩm nang hướng dẫn và bảng giá thiết bị NextFarm."),
        ("Quy mô kho cây trồng", "Tập trung 3 cây chủ lực (Lúa, Cà chua, Dưa chuột) với 185 chunks.", "Giai đoạn 2: Mở rộng sang cây ăn trái/công nghiệp (Sầu riêng, Xoài, Cà phê...) quy mô 2.000+ chunks.")
    ]

    for idx, r in enumerate(road_data, start=1):
        format_row(t_road.rows[idx], [r[0], r[1], r[2]], is_even=(idx % 2 == 0))
        t_road.rows[idx].cells[0].paragraphs[0].runs[0].font.bold = True

    # Save
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOCX))
    doc.save(str(ROOT_DOCX))
    print(f"[OK] Đã tạo thành công file Word tại:\n- {OUT_DOCX}\n- {ROOT_DOCX}")

if __name__ == "__main__":
    main()
